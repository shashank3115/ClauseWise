"""Text extraction utilities for various document formats."""

import asyncio
import io
import logging
import re
from pathlib import Path
from typing import List

import PyPDF2
import docx

logger = logging.getLogger(__name__)


class TextExtractor:
    """Handles text extraction from various document formats."""
    
    def __init__(self, extraction_timeout: int = 30):
        self.extraction_timeout = extraction_timeout
        self.max_pages = 100  # Reasonable limit for PDF pages
    
    def extract_text(self, file_content: bytes, filename: str) -> str:
        """Synchronous wrapper for extract_text - for backwards compatibility."""
        import asyncio
        
        # Try to get existing event loop, or create a new one
        try:
            loop = asyncio.get_event_loop()
            if loop.is_running():
                # If loop is already running, we need to use asyncio.create_task
                # But this is tricky in sync context, so we'll create new loop
                import concurrent.futures
                with concurrent.futures.ThreadPoolExecutor() as executor:
                    future = executor.submit(asyncio.run, self.extract_text_async(file_content, filename))
                    return future.result()
            else:
                return loop.run_until_complete(self.extract_text_async(file_content, filename))
        except RuntimeError:
            # No event loop exists, create one
            return asyncio.run(self.extract_text_async(file_content, filename))
    
    async def extract_text_async(self, file_content: bytes, filename: str) -> str:
        """Async version - this is the main implementation."""
        file_ext = Path(filename).suffix.lower()
        
        try:
            if file_ext == '.pdf':
                text = await asyncio.wait_for(
                    self._extract_from_pdf(file_content), 
                    timeout=self.extraction_timeout
                )
            elif file_ext == '.docx':
                text = await asyncio.wait_for(
                    self._extract_from_docx(file_content), 
                    timeout=self.extraction_timeout
                )
            elif file_ext == '.txt':
                text = await asyncio.wait_for(
                    self._extract_from_txt(file_content), 
                    timeout=10.0  # Shorter timeout for text files
                )
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
            
            if not text or not text.strip():
                raise ValueError("No text content extracted from file")
            
            return text
            
        except asyncio.TimeoutError:
            raise ValueError(f"Text extraction timeout for {filename}")
        except Exception as e:
            logger.error(f"Text extraction failed for {filename}: {str(e)}")
            raise ValueError(f"Failed to extract text: {str(e)}")
    
    async def _extract_from_txt(self, file_content: bytes) -> str:
        """Extract text from TXT file with encoding detection."""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                return file_content.decode(encoding).strip()
            except UnicodeDecodeError:
                continue
        
        raise ValueError("Unable to decode text file with supported encodings")
    
    async def _extract_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF with enhanced error handling and version compatibility."""
        try:
            pdf_file = io.BytesIO(file_content)
            
            # Handle both old and new PyPDF2 versions
            # Check which version is available by checking attributes
            if hasattr(PyPDF2, 'PdfReader'):
                # New version (3.0+)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                pages = pdf_reader.pages
            else:
                # Old version (2.x and earlier)
                pdf_reader = PyPDF2.PdfFileReader(pdf_file)
                pages = [pdf_reader.getPage(i) for i in range(pdf_reader.numPages)]
            
            if len(pages) == 0:
                raise ValueError("PDF contains no pages")
            
            if len(pages) > self.max_pages:
                raise ValueError(f"PDF too large (>{self.max_pages} pages)")
            
            text_content = []
            failed_pages = 0
            
            for page_num, page in enumerate(pages):
                try:
                    # Handle both old and new methods
                    if hasattr(page, 'extract_text'):
                        page_text = page.extract_text()  # New version
                    else:
                        page_text = page.extractText()   # Old version
                        
                    if page_text and page_text.strip():
                        text_content.append(page_text)
                except Exception as e:
                    failed_pages += 1
                    logger.warning(f"Failed to extract text from PDF page {page_num + 1}: {e}")
                    
                    # If too many pages fail, abort
                    if failed_pages > len(pages) * 0.5:
                        raise ValueError("Too many pages failed to extract text")
            
            if not text_content:
                raise ValueError("No text content found in PDF")
            
            return '\n'.join(text_content)
            
        except Exception as e:
            if "No text content found" in str(e) or "too large" in str(e).lower():
                raise  # Re-raise specific errors
            logger.error(f"PDF extraction failed: {str(e)}")
            raise ValueError("Failed to extract text from PDF")
    
    async def _extract_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX with enhanced error handling."""
        try:
            doc_file = io.BytesIO(file_content)
            doc = docx.Document(doc_file)
            
            text_content = []
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text and paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract from tables
            for table in doc.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    text_content.append(table_text)
            
            # Extract from headers and footers
            for section in doc.sections:
                header_text = self._extract_header_footer_text(section.header)
                footer_text = self._extract_header_footer_text(section.footer)
                
                if header_text:
                    text_content.insert(0, header_text)  # Add to beginning
                if footer_text:
                    text_content.append(footer_text)  # Add to end
            
            if not text_content:
                raise ValueError("No text content found in DOCX")
            
            return '\n'.join(text_content)
            
        except Exception as e:
            if "No text content found" in str(e):
                raise  # Re-raise specific errors
            logger.error(f"DOCX extraction failed: {str(e)}")
            raise ValueError("Failed to extract text from Word document")
    
    def _extract_table_text(self, table) -> str:
        """Extract text from a DOCX table."""
        table_content = []
        
        for row in table.rows:
            row_content = []
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    row_content.append(cell.text.strip())
            
            if row_content:
                table_content.append(' | '.join(row_content))
        
        return '\n'.join(table_content) if table_content else ''
    
    def _extract_header_footer_text(self, header_or_footer) -> str:
        """Extract text from header or footer."""
        if not header_or_footer:
            return ''
        
        text_parts = []
        for paragraph in header_or_footer.paragraphs:
            if paragraph.text and paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        return '\n'.join(text_parts) if text_parts else ''


class DocumentMetadataExtractor:
    """Extracts metadata and insights from document content."""
    
    def __init__(self):
        self.min_text_for_analysis = 50
    
    async def extract_metadata(self, text: str, filename: str, file_size: int) -> dict:
        """Extract comprehensive metadata from document text."""
        if not text or len(text.strip()) < self.min_text_for_analysis:
            return {
                'filename': filename,
                'error': 'Insufficient text content for analysis'
            }
        
        try:
            metadata = {
                'filename': filename,
                'file_size': file_size,
                'word_count': len(text.split()),
                'character_count': len(text),
                'paragraph_count': len([p for p in text.split('\n\n') if p.strip()]),
                'detected_language': await self._detect_language(text),
                'contract_type': await self._detect_contract_type(text),
                'parties': await self._extract_parties(text),
                'dates': await self._extract_dates(text),
                'jurisdiction_hints': await self._detect_jurisdiction_hints(text),
                'text_quality_score': await self._assess_text_quality(text),
                'complexity_indicators': await self._analyze_complexity(text)
            }
            
            return metadata
            
        except Exception as e:
            logger.error(f"Metadata extraction failed for {filename}: {str(e)}")
            return {
                'filename': filename,
                'error': f'Metadata extraction failed: {str(e)}'
            }
    
    async def _detect_language(self, text: str) -> str:
        """Enhanced language detection with confidence scoring."""
        if not text:
            return 'unknown'
        
        text_lower = text.lower()
        text_words = text_lower.split()
        
        if len(text_words) < 10:
            return 'en'  # Default for short texts
        
        # Language indicators with weights
        language_indicators = {
            'en': {
                'terms': ['contract', 'agreement', 'whereas', 'party', 'clause', 'terms', 'conditions', 'shall', 'hereby'],
                'patterns': [r'\bwhereas\b', r'\btherefore\b', r'\bshall\b']
            },
            'ms': {
                'terms': ['kontrak', 'perjanjian', 'pihak', 'terma', 'fasal', 'syarat', 'adalah'],
                'patterns': [r'\badalah\b', r'\byang\b', r'\bini\b']
            }
        }
        
        scores = {}
        for lang, indicators in language_indicators.items():
            score = 0
            
            # Term matching
            for term in indicators['terms']:
                score += text_lower.count(term) * 2
            
            # Pattern matching
            for pattern in indicators['patterns']:
                score += len(re.findall(pattern, text_lower))
            
            scores[lang] = score
        
        detected_lang = max(scores, key=scores.get) if scores else 'en'
        return detected_lang if scores[detected_lang] > 0 else 'en'
    
    async def _detect_contract_type(self, text: str) -> str:
        """Enhanced contract type detection using confidence scoring."""
        # Implementation moved from DocumentProcessorService
        # This is a simplified version - you can expand with the full logic
        if not text or len(text) < 50:
            return 'general'
        
        text_lower = text.lower()
        
        # Simple keyword matching for common contract types
        type_keywords = {
            'employment': ['employment', 'employee', 'salary', 'job title'],
            'nda': ['non-disclosure', 'confidentiality', 'proprietary'],
            'service': ['service agreement', 'consultant', 'contractor'],
            'data_processing': ['data processing', 'gdpr', 'personal data']
        }
        
        scores = {}
        for contract_type, keywords in type_keywords.items():
            score = sum(text_lower.count(keyword) for keyword in keywords)
            scores[contract_type] = score
        
        if not any(scores.values()):
            return 'general'
        
        return max(scores, key=scores.get)
    
    async def _extract_parties(self, text: str) -> List[str]:
        """Extract contract parties with validation."""
        if not text:
            return []
        
        parties = []
        
        # Pattern for party extraction
        patterns = [
            r'between\s+([A-Z][A-Za-z\s&.,()-]+?)\s+and\s+([A-Z][A-Za-z\s&.,()-]+?)(?:\s|,|\n)',
            r'Party\s+[A-Z]:\s*([A-Z][A-Za-z\s&.,()-]+?)(?:\n|,|;)'
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
            if matches:
                if isinstance(matches[0], tuple):
                    parties.extend([p.strip(' .,()') for match in matches for p in match])
                else:
                    parties.extend([p.strip(' .,()') for p in matches])
        
        # Clean and validate
        cleaned_parties = []
        for party in parties:
            if party and 2 < len(party) < 100:
                if not any(noise in party.lower() for noise in ['page', 'section', 'article']):
                    cleaned_parties.append(party)
        
        return list(dict.fromkeys(cleaned_parties))[:5]  # Remove duplicates, limit to 5
    
    async def _extract_dates(self, text: str) -> List[str]:
        """Extract dates with validation."""
        if not text:
            return []
        
        date_patterns = [
            r'\b\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4}\b',
            r'\b\d{1,2}\s+(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{2,4}\b',
            r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{2,4}\b'
        ]
        
        dates = set()
        for pattern in date_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            dates.update(matches)
        
        return sorted(list(dates))[:10]  # Limit to 10 dates
    
    async def _detect_jurisdiction_hints(self, text: str) -> List[str]:
        """Detect jurisdiction hints."""
        if not text:
            return []
        
        text_lower = text.lower()
        
        # Simplified jurisdiction detection
        jurisdiction_keywords = {
            'MY': ['malaysia', 'kuala lumpur', 'pdpa', 'employment act 1955'],
            'SG': ['singapore', 'pdpa singapore', 'singapore law'],
            'US': ['united states', 'california', 'ccpa', 'federal law'],
            'EU': ['gdpr', 'european union', 'data protection']
        }
        
        detected = []
        for jurisdiction, keywords in jurisdiction_keywords.items():
            if any(keyword in text_lower for keyword in keywords):
                detected.append(jurisdiction)
        
        return detected
    
    async def _assess_text_quality(self, text: str) -> float:
        """Assess text quality (0.0 to 1.0)."""
        if not text:
            return 0.0
        
        score = 0.0
        
        # Length score
        length_score = min(len(text) / 10000, 1.0) if len(text) > 100 else len(text) / 100
        score += length_score * 0.4
        
        # Word ratio score
        words = text.split()
        avg_word_length = sum(len(word) for word in words) / len(words) if words else 0
        word_score = min(avg_word_length / 6, 1.0) if avg_word_length > 2 else 0
        score += word_score * 0.3
        
        # Legal content indicators
        legal_terms = ['contract', 'agreement', 'party', 'clause', 'terms', 'conditions']
        legal_count = sum(1 for term in legal_terms if term.lower() in text.lower())
        legal_score = min(legal_count / 6, 1.0)
        score += legal_score * 0.3
        
        return min(score, 1.0)
    
    async def _analyze_complexity(self, text: str) -> dict:
        """Analyze document complexity indicators."""
        words = text.split()
        sentences = re.split(r'[.!?]+', text)
        
        return {
            'avg_sentence_length': len(words) / len(sentences) if sentences else 0,
            'long_words_ratio': len([w for w in words if len(w) > 6]) / len(words) if words else 0,
            'punctuation_density': len(re.findall(r'[.,;:!?]', text)) / len(text) if text else 0
        }
